# Imports for file handling, NLP, PDF/Word processing, API requests, and temporary file creation
import requests
import PyPDF2
import docx
import spacy
import nltk
import re
from nltk.corpus import stopwords
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit
from io import BytesIO
import os
from openai import OpenAI
from dotenv import load_dotenv
load_dotenv()


# Download necessary NLTK corpora
nltk.download('punkt_tab')
nltk.download('stopwords')

# Set your Perplexity API key here

PERPLEXITY_API_KEY = os.getenv("PERPLEXITY_API_KEY")
if not PERPLEXITY_API_KEY:
    raise ValueError("Perplexity API key not found. Please set the PERPLEXITY_API_KEY environment variable.")

load_dotenv()
# client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
# Load spaCy's English language model
nlp = spacy.load("en_core_web_sm")

# Define stopwords set for keyword filtering
stop_words = set(stopwords.words('english'))

# -----------------------------------------------
# TEXT EXTRACTION FUNCTIONS
# -----------------------------------------------

def extract_text(resume_file):
    """
    Extract text from a resume file (PDF or DOCX).
    """
    if resume_file.name.endswith(".pdf"):
        reader = PyPDF2.PdfReader(resume_file)
        return "".join([page.extract_text() for page in reader.pages])
    elif resume_file.name.endswith(".docx"):
        doc = docx.Document(resume_file)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        return "Unsupported file type."

# -----------------------------------------------
# KEYWORD & SKILL MATCHING (ATS)
# -----------------------------------------------

def extract_keywords(text, top_n=10):
    """
    Tokenize and return top N frequent keywords from given text.
    Filters out stopwords and non-alphabetic tokens.
    """
    words = nltk.word_tokenize(text)
    words = [w.lower() for w in words if w.isalpha() and w.lower() not in stop_words]
    freq_dist = nltk.FreqDist(words)
    return [word for word, freq in freq_dist.most_common(top_n)]

def ats_keyword_check(resume_text, job_desc):
    """
    Compare resume and job description to identify keyword overlap for ATS scoring.
    """
    resume_keywords = set(extract_keywords(resume_text, 50))
    jd_keywords = set(extract_keywords(job_desc, 50))
    matching = resume_keywords & jd_keywords
    coverage = len(matching) / (len(jd_keywords) or 1)
    suggestions = jd_keywords - resume_keywords
    return {
        "matching_keywords": list(matching),
        "missing_keywords": list(suggestions),
        "coverage_percent": round(coverage * 100, 2)
    }

# -----------------------------------------------
# FILE EXPORTING FUNCTIONS
# -----------------------------------------------


import re

def export_to_docx(text):
    """Exports given text into a DOCX file, preserving basic formatting like **bold**."""
    buffer = BytesIO()
    doc = Document()

    bold_pattern = re.compile(r'\*\*(.*?)\*\*')  # Matches **bold**

    for line in text.split('\n'):
        para = doc.add_paragraph()
        if line.startswith("[Score & Feedback]"):
            run = para.add_run("Score & Feedback:\n")
            run.bold = True
            continue

        # If line contains **bold** segments, parse them
        pos = 0
        for match in bold_pattern.finditer(line):
            start, end = match.span()
            if start > pos:
                para.add_run(line[pos:start])
            bold_run = para.add_run(match.group(1))
            bold_run.bold = True
            pos = end
        if pos < len(line):
            para.add_run(line[pos:])

    doc.save(buffer)
    buffer.seek(0)
    return buffer


def export_to_pdf(text):
    """Exports resume text to a formatted PDF file, with basic bold handling."""
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    x_margin = 40
    max_width = width - 2 * x_margin
    y = height - 40
    font_size = 11

    normal_font = "Helvetica"
    bold_font = "Helvetica-Bold"

    line_spacing = 10
    section_spacing = 15

    def draw_wrapped_line(parts):
        nonlocal y
        x = x_margin

        for part, is_bold in parts:
            font = bold_font if is_bold else normal_font
            c.setFont(font, font_size)

            # Split part into words and wrap manually
            words = part.split()
            for word in words:
                word_width = c.stringWidth(word + ' ', font, font_size)
                if x + word_width > width - x_margin:
                    y -= line_spacing
                    x = x_margin
                    if y < 40:
                        c.showPage()
                        y = height - 40
                        c.setFont(font, font_size)
                c.drawString(x, y, word + ' ')
                x += word_width
        y -= line_spacing
        if y < 40:
            c.showPage()
            y = height - 40

    bold_pattern = re.compile(r'\*\*(.*?)\*\*')

    for line in text.strip().split('\n'):
        line = line.strip()
        if not line:
            y -= 10
            continue

        if line.startswith("[Score & Feedback]"):
            c.setFont(bold_font, font_size)
            c.drawString(x_margin, y, "Score & Feedback:")
            y -= line_spacing
            continue

        parts = []
        last = 0
        for match in bold_pattern.finditer(line):
            if match.start() > last:
                parts.append((line[last:match.start()], False))
            parts.append((match.group(1), True))
            last = match.end()
        if last < len(line):
            parts.append((line[last:], False))

        draw_wrapped_line(parts)

    c.save()
    buffer.seek(0)
    return buffer

# -----------------------------------------------
# PERPLEXITY LLM INTEGRATION
# -----------------------------------------------

def call_perplexity(prompt):
    """
    Call Perplexity LLM API with the given prompt.
    Returns the response as plain text.
    """
    url = "https://api.perplexity.ai/chat/completions"
    headers = {
        "Authorization": f"Bearer {PERPLEXITY_API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "sonar-pro",
        "messages": [{"role": "user", "content": prompt}]
    }
    response = requests.post(url, headers=headers, json=data)
    response.raise_for_status()
    return response.json()['choices'][0]['message']['content']
    # if you want to call open ai 
    # response = client.chat.completions.create(
    #     model="gpt-3.5-turbo",
    #     messages=[
    #         {"role": "user", "content": "Hello!"}
    #     ]
    # )

    # print(response.choices[0].message.content)

# -----------------------------------------------
# SKILL EXTRACTION / PARSING
# -----------------------------------------------

def fetch_dynamic_skillset_from_perplexity(job_desc):
    """
    Uses LLM to extract up to 50 relevant skills from the job description.
    Returns a set of normalized skill terms.
    """
    prompt = (
        "Extract a list of up to 50 relevant technical and soft skills from the following job description. "
        "Only return a comma-separated list of skill names, no explanations:\n\n"
        + job_desc
    )
    skills_raw = call_perplexity(prompt)
    skills = re.split(r',\s*|\n', skills_raw)
    return set(skill.strip().lower() for skill in skills if skill.strip())

def extract_skills(text, skillset):
    """
    Extract known skills from resume text using a given skillset.
    Returns a set of matched skills.
    """
    doc = nlp(text)
    return set(token.text.lower() for token in doc if token.text.lower() in skillset)

# -----------------------------------------------
# STRUCTURED SECTION PARSING VIA LLM
# -----------------------------------------------

def parse_sections_with_llm(resume_text):
    """
    Parse unstructured resume text into structured sections using LLM.
    Supports custom or non-standard section detection.
    """
    prompt = (
        "You are an expert resume parser. Given the following unstructured resume text, extract and organize it into a clean structured format. "
        "Detect all relevant sections, even if they are not standard (like Certifications, Projects, Publications, Languages, etc). "
        "Return each section with a clear heading like this:\n\n"
        "=== Section Name ===\n"
        "Section content...\n\n"
        "Make sure the output is complete and readable.\n\n"
        f"Resume:\n{resume_text}"
    )
    return call_perplexity(prompt)

# -----------------------------------------------
# SECTION SCORING
# -----------------------------------------------

def score_section_with_llm(section_name, section_content):
    """
    Use LLM to score a resume section (1-10) with feedback based on clarity, impact, and relevance.
    """
    prompt = (
        f"Evaluate the following '{section_name}' section of a resume. "
        "Score it from 1 to 10 based on clarity, impact, and relevance to a typical job description. "
        "Also give 1-2 sentences of constructive feedback.\n\n"
        f"Section Content:\n{section_content}"
    )
    return call_perplexity(prompt)

def score_all_sections(parsed_resume):
    """
    Apply scoring to each section in a structured resume.
    Returns a dictionary of section names to score/feedback.
    """
    sections = re.findall(r'=== (.*?) ===\n(.*?)(?=(?:===|\Z))', parsed_resume, re.DOTALL)
    section_scores = {}
    for title, content in sections:
        result = score_section_with_llm(title.strip(), content.strip())
        section_scores[title.strip()] = result
    return section_scores

def regenerate_section_with_llm(section_name, current_content):
    """
    Regenerates a specific section of the resume using the LLM,
    improving clarity, formatting, and alignment with typical job expectations.
    """
    prompt = (
        f"You are an expert resume writer. Rewrite the '{section_name}' section below to improve clarity, impact, and alignment with industry best practices. "
        "Keep it concise and relevant. Return only the improved content without section headers or additional commentary.\n\n"
        f"{section_name} Section:\n{current_content}"
    )
    return call_perplexity(prompt)
# -----------------------------------------------
# RESUME OPTIMIZATION
# -----------------------------------------------

def optimize_resume(resume_text, job_desc, skillset):
    """
    Use LLM to tailor the resume based on the job description and detected skills.
    Returns optimized text and skill overlap stats.
    """
    prompt = (
        "You are an expert resume optimizer. Given the following resume and job description, tailor the resume towards the job, "
        "add important missing keywords and skills, and suggest ATS improvements. Highlight additions and changes.\n\n"
        f"Relevant Skills: {', '.join(skillset)}\n\n"
        f"Resume:\n{resume_text}\n\nJob Description:\n{job_desc}"
    )
    optimized_resume = call_perplexity(prompt)
    resume_skills = extract_skills(resume_text.lower(), skillset)
    jd_skills = extract_skills(job_desc.lower(), skillset)
    return optimized_resume, resume_skills, jd_skills
